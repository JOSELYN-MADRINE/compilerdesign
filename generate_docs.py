"""
Generate Compiler Project Documentation as a .docx Word file.
Run:  python3 generate_docs.py
Output: Compiler_Documentation.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────── colour palette ──────────────────────
DARK_BLUE  = RGBColor(0x1A, 0x36, 0x5C)   # headings
MID_BLUE   = RGBColor(0x27, 0x5E, 0x8E)   # subheadings
ACCENT     = RGBColor(0x2E, 0x86, 0xC1)   # highlights
CODE_BG    = RGBColor(0xF2, 0xF3, 0xF4)   # table / code rows
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x00, 0x00, 0x00)
LIGHT_GREY = RGBColor(0xE8, 0xEA, 0xED)
DARK_GREY  = RGBColor(0x44, 0x44, 0x44)
GREEN      = RGBColor(0x1E, 0x8B, 0x4C)
RED        = RGBColor(0xC0, 0x39, 0x2B)


# ─────────────────────────────────────── helpers ─────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table-cell background via XML shading."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_row_bg(row, hex_color: str):
    for cell in row.cells:
        set_cell_bg(cell, hex_color)


def add_horizontal_rule(doc: Document):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '2E86C1')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = DARK_BLUE if level == 1 else MID_BLUE
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    return p


def body(doc: Document, text: str, bold=False, italic=False,
         color: RGBColor = None, size: int = 11, align=WD_ALIGN_PARAGRAPH.LEFT):
    p   = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = color or DARK_GREY
    return p


def bullet(doc: Document, text: str, level: int = 0, bold_part: str = ""):
    p    = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    if bold_part:
        r = p.add_run(bold_part)
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = DARK_GREY
        rest = text[len(bold_part):]
        r2 = p.add_run(rest)
        r2.font.size = Pt(11)
        r2.font.color.rgb = DARK_GREY
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.color.rgb = DARK_GREY
    return p


def code_block(doc: Document, code: str):
    """Render a grey monospace code block."""
    for line in code.strip().split('\n'):
        p   = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.4)
        p.paragraph_format.right_indent = Inches(0.4)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(line if line else ' ')
        run.font.name  = 'Courier New'
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0x17, 0x20, 0x2A)
        # light background via shading
        pPr  = p._p.get_or_add_pPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F2F3F4')
        pPr.append(shd)


def make_table(doc: Document, rows: list, headers: tuple,
               col_widths: tuple):
    """Generic N-column table builder."""
    ncols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = tbl.rows[0]
    set_row_bg(hdr, '1A365C')
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.width = Inches(col_widths[i])
        p   = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.bold       = True
        run.font.size       = Pt(10)
        run.font.color.rgb  = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for idx, row_data in enumerate(rows):
        row = tbl.rows[idx + 1]
        bg  = 'FFFFFF' if idx % 2 == 0 else 'EAF2FB'
        set_row_bg(row, bg)
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.width = Inches(col_widths[ci])
            p    = cell.paragraphs[0]
            run  = p.add_run(str(val))
            run.font.size      = Pt(10)
            run.font.color.rgb = DARK_GREY
            if ci == 0:
                run.font.bold = True

    return tbl


def two_col_table(doc: Document, rows: list, headers=('Item', 'Description'),
                  col_widths=(2.2, 4.0)):
    return make_table(doc, rows, headers, col_widths)


# ═════════════════════════════════════════════════════════════════ BUILD DOC ══

def build():
    doc = Document()

    # ── page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ══════════════════════════════════════════════════════ TITLE PAGE ════════
    doc.add_paragraph()   # top spacer

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PYTHON COMPILER PROJECT")
    r.font.size  = Pt(28)
    r.font.bold  = True
    r.font.color.rgb = DARK_BLUE

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Technical Documentation")
    rs.font.size  = Pt(16)
    rs.font.color.rgb = MID_BLUE
    rs.font.italic = True

    add_horizontal_rule(doc)

    meta_items = [
        ("Project",   "Python Compiler — Lexical, Syntax & Semantic Analysis"),
        ("Language",  "Python 3.x"),
        ("Interface", "Tkinter GUI"),
        ("Author",    "Compiler Design Project"),
        ("Date",      datetime.date.today().strftime("%d %B %Y")),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rb = p.add_run(f"{label}:  ")
        rb.font.bold  = True
        rb.font.size  = Pt(11)
        rb.font.color.rgb = DARK_BLUE
        rv = p.add_run(value)
        rv.font.size  = Pt(11)
        rv.font.color.rgb = DARK_GREY

    doc.add_page_break()

    # ══════════════════════════════════════════════════ TABLE OF CONTENTS ═════
    heading(doc, "Table of Contents", 1)
    toc_entries = [
        ("1.", "Project Overview"),
        ("2.", "System Requirements"),
        ("3.", "Project Structure"),
        ("4.", "Supported Language"),
        ("5.", "Phase 1 – Lexical Analysis"),
        ("6.", "Phase 2 – Syntax Analysis"),
        ("  6.1", "Top-Down Parser (Recursive Descent)"),
        ("  6.2", "Bottom-Up Parser (Shift-Reduce)"),
        ("7.", "Phase 3 – Semantic Analysis"),
        ("8.", "Abstract Syntax Tree (AST)"),
        ("9.", "Symbol Table"),
        ("10.", "GUI Frontend"),
        ("11.", "Error Handling & Recovery"),
        ("12.", "Sample Programs"),
        ("13.", "Running the Compiler"),
        ("14.", "Limitations & Future Work"),
    ]
    for num, title in toc_entries:
        p   = doc.add_paragraph()
        tab = p.add_run(f"  {num:<8}{title}")
        tab.font.size = Pt(11)
        tab.font.color.rgb = DARK_GREY
        if not num.startswith(" "):
            tab.font.bold = True

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════ SECTION 1 ═
    heading(doc, "1.  Project Overview", 1)
    add_horizontal_rule(doc)
    body(doc, (
        "This project implements a fully functional compiler written in Python that "
        "processes a custom high-level programming language through three classical "
        "compilation phases: Lexical Analysis, Syntax Analysis, and Semantic Analysis. "
        "The compiler is accompanied by a graphical IDE built with Tkinter that "
        "allows users to write, compile, and inspect every phase of compilation "
        "interactively."
    ))
    doc.add_paragraph()
    body(doc, "Key features at a glance:", bold=True)
    bullet(doc, "Three-phase compiler pipeline (Lexer → Parser → Semantic Analyser)")
    bullet(doc, "Two interchangeable parsers: Top-Down (LL) and Bottom-Up (LR/Shift-Reduce)")
    bullet(doc, "Full Abstract Syntax Tree (AST) construction and pretty-printing")
    bullet(doc, "Scoped symbol table with type tracking")
    bullet(doc, "Detailed per-phase output visible in the GUI")
    bullet(doc, "Syntax-highlighted source editor with line numbers")
    bullet(doc, "Open / Save file support")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════ SECTION 2 ═
    heading(doc, "2.  System Requirements", 1)
    add_horizontal_rule(doc)
    two_col_table(doc, [
        ("Python",    "Python 3.10 or later"),
        ("Tkinter",   "Included with the standard Python distribution"),
        ("OS",        "Windows 10/11, macOS 12+, or any Linux desktop (Ubuntu, Fedora, etc.)"),
        ("RAM",       "128 MB minimum (compiler is lightweight)"),
        ("Disk",      "< 1 MB for all source files"),
        ("Libraries", "No third-party Python packages required at runtime"),
    ], headers=("Requirement", "Details"))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════ SECTION 3 ═
    heading(doc, "3.  Project Structure", 1)
    add_horizontal_rule(doc)
    body(doc, "The project is organised as eight Python source files, each responsible for a single concern:")
    doc.add_paragraph()
    two_col_table(doc, [
        ("tokens.py",    "Defines the TokenType enumeration and the Token dataclass used throughout the compiler."),
        ("lexer.py",     "Lexical analyser (Phase 1). Scans source text character-by-character and emits a stream of Token objects."),
        ("ast_nodes.py", "Dataclass definitions for every node type in the Abstract Syntax Tree."),
        ("parser_td.py", "Top-Down (Recursive Descent) parser (Phase 2a). Implements a hand-written LL parser that builds the AST."),
        ("parser_bu.py", "Bottom-Up (Shift-Reduce) parser (Phase 2b). Implements an LR-style parser that explicitly logs every SHIFT and REDUCE action."),
        ("semantic.py",  "Semantic analyser (Phase 3). Walks the AST to enforce type correctness and scoping rules."),
        ("compiler.py",  "Orchestrator that chains all three phases and exposes a single compile() method to the GUI."),
        ("gui.py",       "Tkinter IDE frontend. Main entry point — run this file to start the application."),
    ], headers=("File", "Responsibility"), col_widths=(1.6, 4.6))

    body(doc, "\nDependency flow:")
    code_block(doc, """\
gui.py
  └── compiler.py
        ├── lexer.py         (tokens.py)
        ├── parser_td.py     (tokens.py, ast_nodes.py)
        ├── parser_bu.py     (tokens.py, ast_nodes.py)
        └── semantic.py      (ast_nodes.py)""")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════ SECTION 4 ═
    heading(doc, "4.  Supported Language", 1)
    add_horizontal_rule(doc)
    body(doc, "The compiler processes a statically-typed, imperative mini-language with the following constructs:")
    doc.add_paragraph()

    heading(doc, "4.1  Data Types", 2)
    two_col_table(doc, [
        ("int",    "32-bit integer literal, e.g.  42"),
        ("float",  "Floating-point literal, e.g.  3.14"),
        ("string", 'Double-quoted text, e.g.  "Hello"'),
        ("bool",   "Boolean literal:  true  or  false"),
    ], headers=("Type", "Description"), col_widths=(1.5, 4.7))

    doc.add_paragraph()
    heading(doc, "4.2  Statements", 2)
    two_col_table(doc, [
        ("Declaration",  "int x = 5;  /  float y;"),
        ("Assignment",   "x = x + 1;"),
        ("If / Else",    "if (cond) { … } else { … }"),
        ("While loop",   "while (cond) { … }"),
        ("Print",        "print(expr);"),
        ("Block",        "{ statement* }  — creates a new scope"),
        ("Comment",      "// single-line comment"),
    ], headers=("Construct", "Syntax Example"), col_widths=(1.8, 4.4))

    doc.add_paragraph()
    heading(doc, "4.3  Operators", 2)
    two_col_table(doc, [
        ("Arithmetic",   "+ - * / %"),
        ("Comparison",   "== != < > <= >="),
        ("Logical",      "&& || !"),
        ("Assignment",   "="),
    ], headers=("Category", "Operators"), col_widths=(2.0, 4.2))

    doc.add_paragraph()
    heading(doc, "4.4  Grammar (BNF)", 2)
    body(doc, "The complete grammar used by both parsers:")
    code_block(doc, """\
program     →  stmt_list  EOF
stmt_list   →  stmt*
stmt        →  decl_stmt | assign_stmt | if_stmt | while_stmt | print_stmt | block
decl_stmt   →  type  ID  [ '='  expr ]  ';'
assign_stmt →  ID  '='  expr  ';'
if_stmt     →  'if'  '('  expr  ')'  block  [ 'else'  block ]
while_stmt  →  'while'  '('  expr  ')'  block
print_stmt  →  'print'  '('  expr  ')'  ';'
block       →  '{'  stmt_list  '}'

expr        →  or_expr
or_expr     →  and_expr  ( '||'  and_expr )*
and_expr    →  eq_expr   ( '&&'  eq_expr  )*
eq_expr     →  rel_expr  ( ('=='|'!=')  rel_expr )*
rel_expr    →  add_expr  ( ('<'|'>'|'<='|'>=')  add_expr )*
add_expr    →  mul_expr  ( ('+'|'-')  mul_expr  )*
mul_expr    →  unary     ( ('*'|'/'|'%')  unary   )*
unary       →  ('!'|'-')  unary  |  primary
primary     →  NUMBER | STRING | BOOL | ID | '('  expr  ')'
type        →  'int' | 'float' | 'string' | 'bool'""")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════ SECTION 5 ═
    heading(doc, "5.  Phase 1 – Lexical Analysis", 1)
    add_horizontal_rule(doc)
    body(doc, (
        "The Lexer (lexer.py) is the first phase of compilation. It reads the raw "
        "source code character by character and groups characters into meaningful "
        "units called tokens. Each token carries its type, literal value, and the "
        "source line and column where it was found."
    ))
    doc.add_paragraph()

    heading(doc, "5.1  How it works", 2)
    bullet(doc, "Whitespace and newlines are skipped (they only update the line/column counter).")
    bullet(doc, "// starts a single-line comment; all characters until newline are discarded.")
    bullet(doc, "Digit sequences → INT_LITERAL or FLOAT_LITERAL (if a decimal point is present).")
    bullet(doc, 'Opening " → STRING_LITERAL; escape sequences \\n \\t \\\\ \\" are resolved.')
    bullet(doc, "Alpha / underscore → an identifier; checked against the keyword table → keyword or IDENTIFIER.")
    bullet(doc, "Two-character operators (==, !=, <=, >=, &&, ||) are consumed before single-character ones.")
    bullet(doc, "Unrecognised characters produce a lexer error but scanning continues (error recovery).")

    doc.add_paragraph()
    heading(doc, "5.2  Token categories", 2)
    two_col_table(doc, [
        ("INT_LITERAL",    "Whole number, e.g. 42"),
        ("FLOAT_LITERAL",  "Decimal number, e.g. 3.14"),
        ("STRING_LITERAL", 'Quoted text, e.g. "hello"'),
        ("BOOL_LITERAL",   "true or false"),
        ("IDENTIFIER",     "User-defined name, e.g. counter"),
        ("KW_INT / KW_FLOAT / KW_STRING / KW_BOOL", "Type keywords"),
        ("KW_IF / KW_ELSE / KW_WHILE / KW_PRINT",   "Control keywords"),
        ("PLUS / MINUS / MULTIPLY / DIVIDE / MODULO", "Arithmetic operators"),
        ("EQ / NEQ / LT / GT / LTE / GTE",           "Comparison operators"),
        ("AND / OR / NOT",                            "Logical operators"),
        ("ASSIGN",                                    "= (assignment)"),
        ("LPAREN / RPAREN / LBRACE / RBRACE / SEMICOLON", "Delimiters"),
        ("EOF",                                       "End of input sentinel"),
    ], headers=("Token Type", "Description"), col_widths=(3.0, 3.2))

    doc.add_paragraph()
    heading(doc, "5.3  Example", 2)
    body(doc, "Source code:")
    code_block(doc, "int x = 10 + 5;")
    body(doc, "Tokens produced:")
    code_block(doc, """\
#     Type              Value    Line  Col
1     KW_INT            'int'    1     1
2     IDENTIFIER        'x'      1     5
3     ASSIGN            '='      1     7
4     INT_LITERAL       '10'     1     9
5     PLUS              '+'      1     12
6     INT_LITERAL       '5'      1     14
7     SEMICOLON         ';'      1     15
8     EOF               'EOF'    1     16""")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════ SECTION 6 ═
    heading(doc, "6.  Phase 2 – Syntax Analysis", 1)
    add_horizontal_rule(doc)
    body(doc, (
        "Syntax analysis (parsing) takes the token stream produced by the lexer "
        "and verifies that it conforms to the language grammar.  If it does, an "
        "Abstract Syntax Tree (AST) is constructed.  The compiler ships with two "
        "independent parsers that produce identical ASTs from the same input — "
        "each demonstrating a fundamentally different parsing strategy."
    ))
    doc.add_paragraph()

    # ── 6.1 Top-Down ──────────────────────────────────────────────────────────
    heading(doc, "6.1  Top-Down Parser — Recursive Descent (parser_td.py)", 2)
    body(doc, "Strategy: start from the grammar's start symbol (program) and predict which production to expand based on the current lookahead token, recursively descending until terminals are matched.")
    doc.add_paragraph()
    body(doc, "Characteristics:", bold=True)
    bullet(doc, "LL(1) — reads input left-to-right, leftmost derivation, one token of lookahead.")
    bullet(doc, "Each non-terminal in the grammar corresponds directly to one Python method.")
    bullet(doc, "Easy to read and debug; parse steps logged as  ► rule  and  ✓ reduced.")
    bullet(doc, "Panic-mode error recovery: on a syntax error the parser skips tokens until a synchronisation point (;  or  }) is found.")
    doc.add_paragraph()
    body(doc, "Parse step log excerpt:", bold=True)
    code_block(doc, """\
► program → stmt_list EOF
  stmt → decl_stmt  [int at line 1]
    ✓ primary → INT_LITERAL(10)
    ✓ primary → INT_LITERAL(5)
    ✓ add_expr → expr + expr
    ✓ decl_stmt → int x = <expr> ;""")

    doc.add_paragraph()

    # ── 6.2 Bottom-Up ─────────────────────────────────────────────────────────
    heading(doc, "6.2  Bottom-Up Parser — Shift-Reduce (parser_bu.py)", 2)
    body(doc, "Strategy: scan tokens left-to-right, pushing (SHIFTing) them onto a stack. Whenever the top of the stack matches the right-hand side of a production rule, REDUCE: pop the matched symbols and push the resulting non-terminal. Repeat until the start symbol is on the stack.")
    doc.add_paragraph()
    body(doc, "Characteristics:", bold=True)
    bullet(doc, "LR-style — reads input left-to-right, rightmost derivation in reverse.")
    bullet(doc, "More powerful than LL: can handle left-recursive grammars (expressed here in iterative loops).")
    bullet(doc, "Every SHIFT and REDUCE operation is logged for full traceability in the GUI.")
    bullet(doc, "Same error recovery strategy as the Top-Down parser.")
    doc.add_paragraph()
    body(doc, "Parse step log excerpt:", bold=True)
    code_block(doc, """\
SHIFT   │ KW_INT               'int'
SHIFT   │ IDENTIFIER           'x'
SHIFT   │ ASSIGN               '='
SHIFT   │ INT_LITERAL          '10'
REDUCE  │ primary → INT(10)
SHIFT   │ PLUS                 '+'
SHIFT   │ INT_LITERAL          '5'
REDUCE  │ primary → INT(5)
REDUCE  │ add_expr → add_expr + mul_expr
SHIFT   │ SEMICOLON            ';'
REDUCE  │ decl_stmt → int ID = <expr> ;""")

    doc.add_paragraph()
    body(doc, "Comparison summary:", bold=True)
    make_table(doc, [
        ("Direction",      "Top-down (start → terminals)", "Bottom-up (terminals → start)"),
        ("Derivation",     "Leftmost",                     "Rightmost (reversed)"),
        ("Key operation",  "Predict & expand",             "Shift & reduce"),
        ("Trace style",    "Rule predictions (►/✓)",       "SHIFT / REDUCE actions"),
        ("Power",          "LL(1)",                        "LR-style (more expressive)"),
        ("File",           "parser_td.py",                 "parser_bu.py"),
    ], headers=("Aspect", "Top-Down", "Bottom-Up"), col_widths=(1.8, 2.8, 1.6))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════ SECTION 7 ═
    heading(doc, "7.  Phase 3 – Semantic Analysis", 1)
    add_horizontal_rule(doc)
    body(doc, (
        "After a valid AST is built, the Semantic Analyser (semantic.py) walks the "
        "tree to enforce meaning-level rules that the grammar alone cannot express. "
        "It maintains a hierarchical symbol table that mirrors the block-scope "
        "structure of the source program."
    ))
    doc.add_paragraph()

    heading(doc, "7.1  Checks performed", 2)
    two_col_table(doc, [
        ("Declaration before use", "Every identifier must be declared before it is referenced.  Accessing an undeclared variable raises an error."),
        ("No redeclaration",       "Declaring the same variable twice in the same scope is an error.  Shadowing in inner scopes is allowed."),
        ("Type compatibility",     "The right-hand side of an assignment or initialiser must be compatible with the declared type (int → float widening is permitted)."),
        ("Operator type rules",    "Arithmetic operators require numeric operands; '!' requires bool; string '+' (concatenation) is allowed."),
        ("Condition types",        "The condition of if and while should be bool.  A non-bool condition produces a warning (not a hard error)."),
        ("Comparison types",       "Both operands of a comparison must be the same type (int and float may be mixed)."),
    ], headers=("Check", "Description"), col_widths=(2.2, 4.0))

    doc.add_paragraph()
    heading(doc, "7.2  Scoping rules", 2)
    body(doc, (
        "The analyser creates a new SymbolTable scope each time it enters a block "
        "( { … } ) and destroys it on exit.  Each scope holds a reference to its "
        "parent, so lookup walks up the chain until the global scope is reached or "
        "the variable is not found."
    ))
    code_block(doc, """\
int x = 1;           // global scope, level 0
if (x > 0) {
    int y = 2;       // scope level 1
    x = y + 1;       // OK — x visible from parent scope
}
print(y);            // ERROR — y is out of scope here""")

    doc.add_paragraph()
    heading(doc, "7.3  Analysis log excerpt", 2)
    code_block(doc, """\
Checking: program
  Checking decl: int x  (line 1)
    Defined  x : int
  Checking decl: int x  (line 2)
  [ERROR]   Line 2: Variable 'x' already declared in this scope
  Checking assign: y  (line 3)
  [ERROR]   Line 3: Undefined variable 'y'""")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════ SECTION 8 ═
    heading(doc, "8.  Abstract Syntax Tree (AST)", 1)
    add_horizontal_rule(doc)
    body(doc, (
        "The AST is an in-memory tree of ASTNode objects (defined in ast_nodes.py) "
        "built by both parsers.  It strips away punctuation (semicolons, parentheses, "
        "braces) and retains only the semantic structure of the program. "
        "The compiler.py module provides format_ast() which pretty-prints the tree "
        "as indented text for display in the AST tab of the GUI."
    ))
    doc.add_paragraph()

    heading(doc, "8.1  Node types", 2)
    two_col_table(doc, [
        ("Program",       "Root node; contains a list of top-level statements."),
        ("Block",         "A braced group { … }; creates a scope during semantic analysis."),
        ("DeclStmt",      "Type + name + optional initialiser."),
        ("AssignStmt",    "Name + right-hand expression."),
        ("IfStmt",        "Condition + then-block + optional else-block."),
        ("WhileStmt",     "Condition + body block."),
        ("PrintStmt",     "A single expression to print."),
        ("BinaryExpr",    "Left operand + operator string + right operand."),
        ("UnaryExpr",     "Operator string + single operand."),
        ("NumberLiteral", "Numeric value + is_int flag."),
        ("StringLiteral", "Raw string value."),
        ("BoolLiteral",   "Python bool value."),
        ("Identifier",    "Variable name + source line number."),
    ], headers=("Node Class", "Purpose"), col_widths=(2.0, 4.2))

    doc.add_paragraph()
    heading(doc, "8.2  Example AST output", 2)
    body(doc, 'Source:  int result = x * 2 + 5;')
    code_block(doc, """\
Program
  DeclStmt [int] result
    BinaryExpr ( + )
      BinaryExpr ( * )
        Identifier  x
        Literal [int] 2
      Literal [int] 5""")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════ SECTION 9 ═
    heading(doc, "9.  Symbol Table", 1)
    add_horizontal_rule(doc)
    body(doc, (
        "The SymbolTable is a dictionary of Symbol objects, organised in a linked "
        "list of scopes (each scope stores a reference to its parent). "
        "A Symbol records the variable name, declared type, and source line number."
    ))
    doc.add_paragraph()
    body(doc, "Key operations:", bold=True)
    bullet(doc, "define(sym) — registers a symbol in the current scope.")
    bullet(doc, "lookup(name) — searches the current scope, then walks up the parent chain.")
    bullet(doc, "lookup_local(name) — checks only the current scope (used for redeclaration detection).")
    bullet(doc, "all_symbols() — returns a merged dictionary of all visible symbols, used by the GUI's Symbol Table tab.")
    doc.add_paragraph()
    body(doc, "Example symbol table for the sample program:")
    two_col_table(doc, [
        ("x",       "int     (line 2)"),
        ("pi",      "float   (line 3)"),
        ("msg",     "string  (line 4)"),
        ("flag",    "bool    (line 5)"),
        ("result",  "int     (line 7)"),
        ("i",       "int     (line 14)"),
    ], headers=("Variable", "Type   (declaration line)"), col_widths=(1.6, 4.6))

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════ SECTION 10 ═
    heading(doc, "10.  GUI Frontend (gui.py)", 1)
    add_horizontal_rule(doc)
    body(doc, (
        "The graphical IDE is built entirely with Python's built-in Tkinter library "
        "— no external GUI packages are required.  The window is divided into a "
        "source editor on the left and a tabbed output panel on the right."
    ))
    doc.add_paragraph()

    heading(doc, "10.1  Layout", 2)
    two_col_table(doc, [
        ("Toolbar",         "Run buttons, parser selector (Top-Down / Bottom-Up), Open/Save/Clear/Sample."),
        ("Source Editor",   "Monospace editor with line-number gutter and live syntax highlighting."),
        ("Tokens tab",      "Numbered table of every token: type, value, line, column."),
        ("Parse Steps tab", "Full parse trace — predictions (top-down) or SHIFT/REDUCE log (bottom-up)."),
        ("AST tab",         "Indented pretty-print of the Abstract Syntax Tree."),
        ("Semantic tab",    "Analysis log with errors, warnings, and scope entry/exit markers."),
        ("Symbol Table tab","All declared variables with their types and source lines."),
        ("Summary tab",     "PASS/FAIL per phase plus a consolidated list of all issues."),
        ("Status bar",      "Current operation status and cursor position (Ln / Col)."),
    ], headers=("Component", "Description"), col_widths=(2.0, 4.2))

    doc.add_paragraph()
    heading(doc, "10.2  Syntax highlighting", 2)
    body(doc, "The editor applies live colour tags using Python's re module as the user types:")
    two_col_table(doc, [
        ("Keywords",    "if  else  while  print  true  false  →  purple"),
        ("Types",       "int  float  string  bool  →  blue"),
        ("Numbers",     "integer and float literals  →  orange"),
        ("Strings",     'double-quoted text  →  green'),
        ("Comments",    "// …  →  grey"),
        ("Operators",   "+ - * / = < > ! & |  →  cyan"),
    ], headers=("Category", "Colour"), col_widths=(1.8, 4.4))

    doc.add_paragraph()
    heading(doc, "10.3  Run modes", 2)
    bullet(doc, "Run All Phases — executes Lexer → Parser → Semantic Analyser in sequence.")
    bullet(doc, "Lexer Only — stops after tokenisation and switches to the Tokens tab.")
    bullet(doc, "Parser Only — runs Lexer + selected Parser, shows parse steps and AST.")
    bullet(doc, "Semantic Only — runs all phases and switches to the Semantic tab.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════ SECTION 11 ═
    heading(doc, "11.  Error Handling & Recovery", 1)
    add_horizontal_rule(doc)
    body(doc, "The compiler is designed to report as many errors as possible before stopping.")
    doc.add_paragraph()

    two_col_table(doc, [
        ("Lexer",    "Unknown characters are reported and skipped; scanning continues on the next character."),
        ("Parser",   "On a syntax error, the parser reports the message, then skips tokens until a semicolon or closing brace is found (panic-mode recovery), allowing subsequent statements to be checked."),
        ("Semantic", "Each statement is checked independently; an error in one declaration does not prevent the next from being analysed."),
    ], headers=("Phase", "Recovery Strategy"), col_widths=(1.4, 4.8))

    doc.add_paragraph()
    body(doc, "Error message format:", bold=True)
    code_block(doc, """\
[LEXER]    Line 5, Col 12: Unknown character '@'
[PARSER]   Line 8: Expected SEMICOLON, got IDENTIFIER ('x')
[SEMANTIC] Line 3: Variable 'count' already declared in this scope
[SEMANTIC] Line 7: Cannot assign string to int variable 'x'
[SEMANTIC] Line 9: Undefined variable 'total'""")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════ SECTION 12 ═
    heading(doc, "12.  Sample Programs", 1)
    add_horizontal_rule(doc)

    heading(doc, "12.1  Variables and arithmetic", 2)
    code_block(doc, """\
int    a = 10;
int    b = 3;
float  ratio = 3.14;
string name  = "Alice";
bool   active = true;

int sum  = a + b;
int diff = a - b;
int prod = a * b;""")

    doc.add_paragraph()
    heading(doc, "12.2  Conditionals", 2)
    code_block(doc, """\
int score = 75;

if (score >= 90) {
    print(score);
} else {
    int penalty = 100 - score;
    print(penalty);
}""")

    doc.add_paragraph()
    heading(doc, "12.3  Loops", 2)
    code_block(doc, """\
int i = 1;
int total = 0;

while (i <= 5) {
    total = total + i;
    i = i + 1;
}
print(total);""")

    doc.add_paragraph()
    heading(doc, "12.4  Nested blocks and scoping", 2)
    code_block(doc, """\
int x = 10;

if (x > 0) {
    int y = x * 2;     // y lives only in this block
    print(y);
    if (y > 15) {
        bool big = true;
        print(x);
    }
}
// y and big are not accessible here""")

    doc.add_paragraph()
    heading(doc, "12.5  Semantic error examples", 2)
    code_block(doc, """\
int x = 5;
int x = 10;          // ERROR: redeclaration in same scope
string s = x;        // ERROR: cannot assign int to string
print(undefined);    // ERROR: undefined variable
bool b = 1 + 2;      // ERROR: cannot assign int to bool""")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════ SECTION 13 ═
    heading(doc, "13.  Running the Compiler", 1)
    add_horizontal_rule(doc)

    heading(doc, "13.1  Starting the GUI", 2)
    body(doc, "Open a terminal, navigate to the project folder, and run:")
    code_block(doc, "python3 gui.py")
    body(doc, "The IDE window will open with a sample program pre-loaded.")

    doc.add_paragraph()
    heading(doc, "13.2  Using the compiler step-by-step", 2)
    for n, step in enumerate([
        "Type or paste source code in the left editor panel.",
        "Select a parser:  Top-Down (Recursive Descent)  or  Bottom-Up (Shift-Reduce).",
        "Click  Run All Phases  to execute all three compilation phases.",
        "Switch between the output tabs to inspect tokens, parse steps, AST, semantic results, symbol table, and summary.",
        "Use  Lexer Only,  Parser Only,  or  Semantic Only  to run individual phases.",
        "Click  Open  to load a source file, or  Save  to write the editor contents to disk.",
        "Click  Sample  to reload the built-in example program.",
        "Click  Clear  to reset the editor and all output panels.",
    ], 1):
        bullet(doc, f"{n}.  {step}")

    doc.add_paragraph()
    heading(doc, "13.3  Using the compiler programmatically", 2)
    code_block(doc, """\
from compiler import Compiler

c   = Compiler()
res = c.compile(source_code, parser_type='topdown')  # or 'bottomup'

# Phase outputs
print(res['tokens'])        # list of Token objects
print(res['ast_text'])      # pretty-printed AST string
print(res['parse_steps'])   # list of parse log strings
print(res['sem_errors'])    # list of semantic error strings
print(res['symbol_table'])  # dict  name -> "type  (line N)"
print(res['success'])       # True if all phases passed""")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════ SECTION 14 ═
    heading(doc, "14.  Limitations & Future Work", 1)
    add_horizontal_rule(doc)

    heading(doc, "14.1  Current limitations", 2)
    bullet(doc, "No code generation or interpreter — the compiler analyses but does not execute programs.")
    bullet(doc, "Functions / procedures are not supported in this version.")
    bullet(doc, "Arrays and composite data types are not yet implemented.")
    bullet(doc, "Only single-line // comments are supported (no block comments).")
    bullet(doc, "The Bottom-Up parser uses the same recursive structure as the Top-Down parser to keep the implementation clear; a full LR(1) table-driven parser would require a parser-generator.")

    doc.add_paragraph()
    heading(doc, "14.2  Suggested extensions", 2)
    bullet(doc, "Code Generation — emit Python bytecode or LLVM IR to make programs runnable.")
    bullet(doc, "Function declarations — def f(x: int): … with local scope and return type checking.")
    bullet(doc, "Arrays — int arr[10]; with bounds checking in the semantic phase.")
    bullet(doc, "For loops — for (init; cond; update) { … } syntax.")
    bullet(doc, "Full LR(1) / LALR(1) table-driven bottom-up parser.")
    bullet(doc, "Type inference for var declarations.")
    bullet(doc, "Intermediate representation (three-address code or SSA).")
    bullet(doc, "Optimisation passes (constant folding, dead-code elimination).")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════ BACK PAGE ═
    add_horizontal_rule(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Python Compiler Project  ·  Technical Documentation")
    r.font.size = Pt(10)
    r.font.color.rgb = FG_DIM = RGBColor(0x6E, 0x73, 0x8D)
    r.font.italic = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Generated  {datetime.date.today().strftime('%d %B %Y')}")
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x9A, 0x9A, 0x9A)

    # ── save ─────────────────────────────────────────────────────────────────
    out = "Compiler_Documentation.docx"
    doc.save(out)
    print(f"Documentation saved → {out}")


if __name__ == "__main__":
    build()
