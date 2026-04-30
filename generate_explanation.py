"""
Generates: Compiler_Explanation.docx
A complete walk-through of the project and the built-in example program,
from Lexical Analysis through Semantic Analysis.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────── helpers ────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def body(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def code_block(doc, text, font_size=9):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    # grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(font_size)
    return p

def two_col(doc, rows, header1='', header2='', w1=2.5, w2=4.0):
    """Two-column table with optional header row."""
    n_cols = 2
    has_hdr = bool(header1 or header2)
    tbl = doc.add_table(rows=len(rows) + (1 if has_hdr else 0), cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    r = 0
    if has_hdr:
        for ci, txt in enumerate([header1, header2]):
            c = tbl.rows[0].cells[ci]
            c.text = txt
            c.paragraphs[0].runs[0].bold = True
            c.paragraphs[0].runs[0].font.size = Pt(10)
            set_cell_bg(c, '1F4E79')
            c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r = 1

    for i, (a, b) in enumerate(rows):
        tbl.rows[r + i].cells[0].text = str(a)
        tbl.rows[r + i].cells[1].text = str(b)
        bg = 'EBF3FB' if i % 2 == 0 else 'FFFFFF'
        set_cell_bg(tbl.rows[r + i].cells[0], bg)
        set_cell_bg(tbl.rows[r + i].cells[1], bg)
        for ci in range(2):
            for run in tbl.rows[r + i].cells[ci].paragraphs[0].runs:
                run.font.size = Pt(10)

    set_col_width(tbl, 0, w1)
    set_col_width(tbl, 1, w2)
    return tbl

def three_col(doc, rows, headers=('','',''), widths=(1.8, 1.5, 3.2)):
    tbl = doc.add_table(rows=len(rows) + 1, cols=3)
    tbl.style = 'Table Grid'
    for ci, h in enumerate(headers):
        c = tbl.rows[0].cells[ci]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(c, '1F4E79')
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, row in enumerate(rows):
        bg = 'EBF3FB' if i % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            c = tbl.rows[i + 1].cells[ci]
            c.text = str(val)
            set_cell_bg(c, bg)
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(9)
    for ci, w in enumerate(widths):
        set_col_width(tbl, ci, w)
    return tbl

def divider(doc):
    doc.add_paragraph('─' * 70)

# ═══════════════════════════════════════════════════════════════════ main ════

def build():
    doc = Document()

    # ── page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ══════════════════════════════════════════════════ TITLE PAGE ═══════════
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('Python Compiler — Project & Example Walk-Through')
    r.bold = True; r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run('From Lexical Analysis to Semantic Analysis')
    sr.font.size = Pt(14); sr.italic = True
    sr.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()
    doc.add_paragraph()

    # ══════════════════════════════════════════════ 1. PROJECT OVERVIEW ═══════
    heading(doc, '1.  Project Overview', 1)
    body(doc,
        'This project is a four-phase compiler implemented entirely in Python. '
        'It accepts a simple C-like programming language and processes source code '
        'through each classical compiler stage: Lexical Analysis, Syntax Analysis '
        '(with two different parser strategies), Semantic Analysis, and optional '
        'Execution via a tree-walking interpreter. '
        'A web-based IDE (Flask + HTML/CSS/JS) and a desktop GUI (Tkinter) are '
        'provided so users can write, compile, and run programs interactively.')

    doc.add_paragraph()
    body(doc, 'The four compilation phases are:', bold=True)

    two_col(doc, [
        ('Phase 1 – Lexical Analysis',
         'Reads the raw source characters and groups them into meaningful tokens '
         '(keywords, identifiers, numbers, strings, operators, punctuation).'),
        ('Phase 2 – Syntax Analysis',
         'Checks that the token stream follows the grammar of the language and '
         'builds an Abstract Syntax Tree (AST). Two parser strategies are '
         'implemented: Top-Down (Recursive Descent) and Bottom-Up (Shift-Reduce).'),
        ('Phase 3 – Semantic Analysis',
         'Walks the AST and enforces meaning rules: variables must be declared '
         'before use, types must be compatible, no duplicate declarations in the '
         'same scope.'),
        ('Phase 4 – Execution',
         'A tree-walking interpreter evaluates the AST directly, producing output '
         'from print() statements and tracking runtime state in a scope-chained '
         'environment.'),
    ], header1='Phase', header2='Description')

    # ── language features ────────────────────────────────────────────────────
    doc.add_paragraph()
    heading(doc, '1.1  Supported Language Features', 2)
    two_col(doc, [
        ('Data types',            'int, float, string, bool'),
        ('Variable declaration',  'int x = 10;  /  float y;'),
        ('Assignment',            'x = x + 1;'),
        ('Arithmetic operators',  '+  −  *  /  %'),
        ('Comparison operators',  '==  !=  <  >  <=  >='),
        ('Logical operators',     '&&  ||  !'),
        ('Control flow',          'if / else,  while'),
        ('Output',                'print( <expression> );'),
        ('Comments',              '// single-line comments'),
        ('Block scoping',         'Variables declared inside { } are local to that block'),
    ], header1='Feature', header2='Details', w1=2.2, w2=4.3)

    # ── file layout ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    heading(doc, '1.2  Source File Layout', 2)
    two_col(doc, [
        ('tokens.py',    'TokenType enum, KEYWORDS dict, Token dataclass'),
        ('lexer.py',     'Lexer – Phase 1'),
        ('ast_nodes.py', 'All AST node dataclasses'),
        ('parser_td.py', 'TopDownParser – Phase 2 (Recursive Descent)'),
        ('parser_bu.py', 'BottomUpParser – Phase 2 (Shift-Reduce)'),
        ('semantic.py',  'SemanticAnalyzer, SymbolTable – Phase 3'),
        ('interpreter.py','Interpreter – Phase 4'),
        ('compiler.py',  'Compiler orchestrator – runs all four phases'),
        ('app.py',       'Flask web server  (python3 app.py → localhost:5000)'),
        ('gui.py',       'Tkinter desktop IDE'),
    ], header1='File', header2='Role', w1=2.0, w2=4.5)

    doc.add_page_break()

    # ══════════════════════════════════════════════ 2. THE EXAMPLE PROGRAM ════
    heading(doc, '2.  The Built-In Example Program', 1)
    body(doc,
        'When you click the "Example" button in the web IDE, the following '
        'program is loaded. It exercises every language feature: all four data '
        'types, arithmetic, if/else, boolean logic, while loops, nested blocks, '
        'and string variables. The sections that follow trace this exact program '
        'through every compiler phase.')

    doc.add_paragraph()
    example_code = '''\
// ── 1. All four data types ───────────────────────
int    age     = 20;
float  gpa     = 3.75;
string name    = "Alice";
bool   passing = true;

print("=== Student Report ===");
print(name);
print(age);
print(gpa);

// ── 2. Arithmetic operations ─────────────────────
int  a      = 12;
int  b      = 5;
int  sum    = a + b;
int  diff   = a - b;
int  prod   = a * b;
int  modulo = a % b;

print("=== Arithmetic ===");
print(sum);
print(diff);
print(prod);
print(modulo);

// ── 3. if / else branching ───────────────────────
print("=== Grade Check ===");
if (gpa >= 3.5) {
    print("Distinction");
} else {
    print("Pass");
}

// ── 4. Boolean logic ─────────────────────────────
bool honours = passing && gpa >= 3.5;
if (honours) {
    print("Honours student!");
}

// ── 5. while loop (sum 1 to 10) ──────────────────
print("=== Sum 1 to 10 ===");
int total = 0;
int i     = 1;
while (i <= 10) {
    total = total + i;
    i     = i + 1;
}
print(total);

// ── 6. Nested blocks and scoping ─────────────────
print("=== Countdown ===");
int n = 5;
while (n > 0) {
    if (n == 1) {
        print("Go!");
    } else {
        print(n);
    }
    n = n - 1;
}

// ── 7. String variable ────────────────────────────
string farewell = "Compilation complete!";
print(farewell);'''

    code_block(doc, example_code, font_size=8)

    doc.add_paragraph()
    body(doc,
        'When executed (⚡ Run Code), the program produces the following output:',
        italic=True)

    output_lines = [
        '=== Student Report ===', 'Alice', '20', '3.75',
        '=== Arithmetic ===',    '17', '7', '60', '2',
        '=== Grade Check ===',   'Distinction',
        'Honours student!',
        '=== Sum 1 to 10 ===',   '55',
        '=== Countdown ===',     '5', '4', '3', '2', 'Go!',
        'Compilation complete!',
    ]
    code_block(doc, '\n'.join(output_lines), font_size=9)

    doc.add_page_break()

    # ══════════════════════════════════════════════════ 3. LEXICAL ANALYSIS ════
    heading(doc, '3.  Phase 1 – Lexical Analysis', 1)
    body(doc,
        'Lexical analysis (also called scanning or tokenisation) is the first '
        'phase of a compiler. The Lexer class in lexer.py reads the source code '
        'one character at a time and groups characters into tokens — the smallest '
        'meaningful units of the language.')

    doc.add_paragraph()
    heading(doc, '3.1  What Is a Token?', 2)
    body(doc,
        'A token is a pair (type, value) together with its source position '
        '(line, column). The compiler defines 36 token types in the TokenType '
        'enum inside tokens.py, grouped as follows:')
    doc.add_paragraph()
    two_col(doc, [
        ('INT_LITERAL, FLOAT_LITERAL',  'Numeric constants: 20, 3.75'),
        ('STRING_LITERAL',              'Quoted text: "Alice"'),
        ('BOOL_LITERAL',                'true / false'),
        ('IDENTIFIER',                  'User-defined names: age, gpa, name'),
        ('KW_INT, KW_FLOAT, KW_STRING, KW_BOOL', 'Type keywords'),
        ('KW_IF, KW_ELSE, KW_WHILE, KW_PRINT',   'Control-flow keywords'),
        ('PLUS, MINUS, MULTIPLY, DIVIDE, MODULO', 'Arithmetic operators: + − * / %'),
        ('EQ, NEQ, LT, GT, LTE, GTE',            'Comparison: == != < > <= >='),
        ('AND, OR, NOT',                          'Logical: && || !'),
        ('ASSIGN',                                'Assignment: ='),
        ('LPAREN, RPAREN, LBRACE, RBRACE, SEMICOLON', 'Punctuation: ( ) { } ;'),
        ('EOF',                                   'End of input marker'),
    ], header1='Token Type(s)', header2='Meaning / Example', w1=3.0, w2=3.5)

    doc.add_paragraph()
    heading(doc, '3.2  How the Lexer Works', 2)
    body(doc,
        'The Lexer uses a position pointer (pos) that moves forward '
        'character by character. At each position it decides which kind of '
        'token to build by looking at the current character (and one character '
        'ahead for two-character operators):')
    doc.add_paragraph()
    two_col(doc, [
        ('Whitespace / newlines',     'Skipped silently; newlines increment the line counter.'),
        ('//',                         'Start of line comment — advance to end of line, produce no token.'),
        ('Digit (0-9)',                'Call _read_number(): collect digits, check for decimal point → INT_LITERAL or FLOAT_LITERAL.'),
        ('"',                          'Call _read_string(): collect characters until closing ", handle \\n \\t \\\\ \\" escapes.'),
        ('Letter or _',               'Call _read_identifier(): collect alphanumeric/underscore chars, look up in KEYWORDS dict → keyword token or IDENTIFIER.'),
        ('Two-char pairs: == != <= >= && ||', 'Peek one ahead, consume both characters → single two-character token.'),
        ('Single-char: + - * / % < > ! = ( ) { } ;', 'Look up in SINGLE dict → one token.'),
        ('Anything else',              'Log an error ("Unknown character") and skip.'),
    ], header1='Input Pattern', header2='Lexer Action', w1=2.8, w2=3.7)

    doc.add_paragraph()
    heading(doc, '3.3  Tokenising the First Four Lines of the Example', 2)
    body(doc,
        'Below is the exact token stream produced by the Lexer for the first '
        'four variable declarations and the comment on line 2:')
    doc.add_paragraph()

    three_col(doc, [
        ('int',    'KW_INT',        'Reserved keyword → type token'),
        ('age',    'IDENTIFIER',    'Not in KEYWORDS dict → user-defined name'),
        ('=',      'ASSIGN',        'Single-char SINGLE lookup'),
        ('20',     'INT_LITERAL',   '_read_number(): digits only, no decimal'),
        (';',      'SEMICOLON',     'Single-char SINGLE lookup'),
        ('float',  'KW_FLOAT',      'Reserved keyword → type token'),
        ('gpa',    'IDENTIFIER',    'User-defined name'),
        ('=',      'ASSIGN',        ''),
        ('3.75',   'FLOAT_LITERAL', '_read_number(): digits + decimal point + digits'),
        (';',      'SEMICOLON',     ''),
        ('string', 'KW_STRING',     'Reserved keyword'),
        ('name',   'IDENTIFIER',    'User-defined name'),
        ('=',      'ASSIGN',        ''),
        ('"Alice"','STRING_LITERAL','_read_string(): collect A-l-i-c-e between quotes'),
        (';',      'SEMICOLON',     ''),
        ('bool',   'KW_BOOL',       'Reserved keyword'),
        ('passing','IDENTIFIER',    'User-defined name'),
        ('=',      'ASSIGN',        ''),
        ('true',   'BOOL_LITERAL',  'KEYWORDS["true"] = KW_TRUE → normalised to BOOL_LITERAL'),
        (';',      'SEMICOLON',     ''),
    ],
    headers=('Lexeme (raw text)', 'Token Type', 'How the Lexer decided'),
    widths=(1.5, 1.6, 3.4))

    doc.add_paragraph()
    body(doc,
        'Notice that keywords (int, float, string, bool, true, false, if, else, '
        'while, print) are recognised by looking up the collected identifier in '
        'the KEYWORDS dictionary. If found, the appropriate keyword token type '
        'is returned; otherwise the token becomes an IDENTIFIER.',
        italic=True)

    doc.add_paragraph()
    heading(doc, '3.4  Tokenising the while-loop Condition (line 52)', 2)
    body(doc, 'The condition  i <= 10  inside the while statement becomes:')
    doc.add_paragraph()
    three_col(doc, [
        ('i',    'IDENTIFIER',  'User-defined name for loop counter'),
        ('<=',   'LTE',         '_peek() shows "="; consume both chars → LTE'),
        ('10',   'INT_LITERAL', '_read_number(): integer, no decimal'),
    ],
    headers=('Lexeme', 'Token Type', 'Note'),
    widths=(1.5, 1.6, 3.4))

    doc.add_paragraph()
    body(doc,
        'Key insight: the lexer must peek one character ahead before deciding '
        'whether "<" alone (LT) or "<=" (LTE) is the correct token. The same '
        'lookahead applies to >, >=, ==, !=, &&, and ||.')

    doc.add_page_break()

    # ══════════════════════════════════════════════════ 4. SYNTAX ANALYSIS ═════
    heading(doc, '4.  Phase 2 – Syntax Analysis (Parsing)', 1)
    body(doc,
        'The parser takes the flat list of tokens produced by the Lexer and '
        'checks whether they form a valid program according to the language '
        'grammar. If valid, it builds an Abstract Syntax Tree (AST) — a '
        'tree structure that captures the hierarchical meaning of the program '
        'without all the punctuation noise (semicolons, braces, parentheses).')

    doc.add_paragraph()
    heading(doc, '4.1  The Grammar', 2)
    body(doc,
        'The language grammar is defined by 17 productions. Both parsers '
        'implement the same grammar (they just use different strategies):')
    doc.add_paragraph()
    grammar_text = (
        "program    -> stmt_list EOF\n"
        "stmt_list  -> stmt_list stmt  |  (empty)\n"
        "stmt       -> decl_stmt | assign_stmt | if_stmt | while_stmt | print_stmt | block\n"
        "decl_stmt  -> type ID ';'  |  type ID '=' expr ';'\n"
        "assign_stmt-> ID '=' expr ';'\n"
        "if_stmt    -> 'if' '(' expr ')' block  |  'if' '(' expr ')' block 'else' block\n"
        "while_stmt -> 'while' '(' expr ')' block\n"
        "print_stmt -> 'print' '(' expr ')' ';'\n"
        "block      -> '{' stmt_list '}'\n"
        "expr       -> expr '||' and_expr  |  and_expr\n"
        "and_expr   -> and_expr '&&' eq_expr  |  eq_expr\n"
        "eq_expr    -> eq_expr ('=='|'!=') rel_expr  |  rel_expr\n"
        "rel_expr   -> rel_expr ('<'|'>'|'<='|'>=') add_expr  |  add_expr\n"
        "add_expr   -> add_expr ('+'|'-') mul_expr  |  mul_expr\n"
        "mul_expr   -> mul_expr ('*'|'/'|'%') unary  |  unary\n"
        "unary      -> '!' unary  |  '-' unary  |  primary\n"
        "primary    -> NUMBER | STRING | BOOL | ID | '(' expr ')'"
    )
    code_block(doc, grammar_text, font_size=8)

    body(doc,
        'Operator precedence is built into the grammar: expressions at a lower '
        'level in the hierarchy have higher precedence. For example, mul_expr '
        '(*, /, %) is deeper than add_expr (+, -), so multiplication binds '
        'tighter than addition.', italic=True)

    doc.add_paragraph()
    heading(doc, '4.2  Strategy A — Top-Down Parser (Recursive Descent)', 2)
    body(doc,
        'The TopDownParser in parser_td.py implements a predictive recursive-descent '
        'parser. It starts at the top of the grammar (program) and works downward '
        'to match tokens, predicting which production to use by looking at the '
        'current token (one-symbol lookahead).')
    doc.add_paragraph()
    body(doc, 'Characteristics:', bold=True)
    two_col(doc, [
        ('Direction',         'Top → Down (start symbol → leaves)'),
        ('Lookahead',         '1 token (current token determines which rule to apply)'),
        ('Error recovery',    'Panic-mode: on error, skip forward to the next ; or }'),
        ('Log format',        '"► rule_name" when a rule is predicted; "✓ rule_name" when it completes'),
        ('Implementation',    'One Python method per grammar rule; methods call each other recursively'),
    ], header1='Property', header2='Detail')

    doc.add_paragraph()
    body(doc, 'How it parses  int age = 20;  (Step by Step):', bold=True)
    doc.add_paragraph()
    two_col(doc, [
        ('parse_program() called',       '► sees KW_INT → calls parse_stmt_list()'),
        ('parse_stmt_list() called',      '► sees KW_INT (a type keyword) → calls parse_stmt()'),
        ('parse_stmt() called',           '► current token is KW_INT → calls parse_decl_stmt()'),
        ('parse_decl_stmt() called',      '► consume KW_INT (type_tok = "int")'),
        ('',                              '► consume IDENTIFIER (name_tok = "age")'),
        ('',                              '► current token is ASSIGN → consume it'),
        ('parse_expr() → parse_add() → parse_primary()', '► consume INT_LITERAL "20" → return NumberLiteral(20, is_int=True)'),
        ('Back in parse_decl_stmt()',     '► consume SEMICOLON'),
        ('',                              '✓ decl_stmt → int age = <expr> ;'),
        ('Return value',                  'DeclStmt(var_type="int", name="age", initializer=NumberLiteral(20))'),
    ], header1='Parser Action', header2='Detail', w1=3.2, w2=3.3)

    doc.add_paragraph()
    heading(doc, '4.3  Strategy B — Bottom-Up Parser (Shift-Reduce)', 2)
    body(doc,
        'The BottomUpParser in parser_bu.py implements a Shift-Reduce parser '
        '(LR-style). Instead of starting at the top, it starts at the leaves '
        '(the tokens) and works upward, combining symbols into larger structures '
        'until only the start symbol remains.')
    doc.add_paragraph()
    body(doc, 'Two fundamental operations drive the algorithm:', bold=True)
    two_col(doc, [
        ('SHIFT',  'Push the current input token onto the parse stack and advance the input pointer. '
                   'Logged as:  SHIFT  │  TOKEN_TYPE   \'value\''),
        ('REDUCE', 'When the top of the stack matches the right-hand side of a grammar production, '
                   'pop those symbols and push the corresponding non-terminal (AST node). '
                   'Logged as:  REDUCE │  rule_name'),
    ], header1='Operation', header2='Meaning', w1=1.2, w2=5.3)

    doc.add_paragraph()
    body(doc, 'Shift-Reduce trace for  int age = 20;:', bold=True)
    doc.add_paragraph()
    two_col(doc, [
        ('SHIFT   KW_INT   "int"',          'Token "int" pushed onto stack'),
        ('SHIFT   IDENTIFIER   "age"',       'Token "age" pushed onto stack'),
        ('SHIFT   ASSIGN   "="',             'Token "=" pushed onto stack'),
        ('SHIFT   INT_LITERAL   "20"',       'Token "20" pushed onto stack'),
        ('REDUCE  primary → INT(20)',         'NumberLiteral(20) replaces "20" on stack'),
        ('REDUCE  unary → primary',           'Unary wrapper (no operator, passes through)'),
        ('REDUCE  mul_expr → unary',          'mul_expr non-terminal formed'),
        ('REDUCE  add_expr → mul_expr',       'add_expr non-terminal formed'),
        ('REDUCE  ... → add_expr',            '(rel, eq, and, or all pass through similarly)'),
        ('SHIFT   SEMICOLON   ";"',           'Semicolon consumed'),
        ('REDUCE  decl_stmt → int ID = <expr> ;', 'DeclStmt(int, age, NumberLiteral(20)) formed'),
        ('REDUCE  stmt_list → stmt_list stmt','Statement added to program list'),
    ], header1='Action', header2='Effect', w1=3.0, w2=3.5)

    doc.add_paragraph()
    heading(doc, '4.4  Comparing the Two Parsers', 2)
    two_col(doc, [
        ('',               'Top-Down (Recursive Descent)      vs.      Bottom-Up (Shift-Reduce)'),
        ('Starting point', 'Start symbol (program)',              ),
        ('Direction',      'Predict → expand downward',           ),
        ('Token handling', 'Consume token when rule expects it',  ),
        ('Log style',      '► predicted rule / ✓ completed rule', ),
        ('Strength',       'Easy to write; clear error messages', ),
        ('AST produced',   'Identical in both cases',             ),
    ], header1='Aspect', header2='Comparison', w1=2.0, w2=4.5)

    # Redo as a proper comparison table
    doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)
    tbl2 = doc.add_table(rows=8, cols=3)
    tbl2.style = 'Table Grid'
    hdrs = ['Aspect', 'Top-Down (Recursive Descent)', 'Bottom-Up (Shift-Reduce)']
    for ci, h in enumerate(hdrs):
        c = tbl2.rows[0].cells[ci]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(c, '1F4E79')
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    rows_data = [
        ('Starting point', 'Start symbol (program)',              'Leaves (individual tokens)'),
        ('Direction',      'Expand downward',                     'Combine upward'),
        ('Token use',      'Consume when rule expects it',        'SHIFT onto stack; REDUCE when top matches RHS'),
        ('Log style',      '► predicted / ✓ completed',           'SHIFT │ TOKEN … / REDUCE │ rule …'),
        ('Grammar visible?','No (implicit in call stack)',        'Yes — displayed at parse start'),
        ('Strength',       'Easy to read; clear error messages',  'Explicit shift/reduce trace; shows LR concept'),
        ('AST output',     'Identical program tree',              'Identical program tree'),
    ]
    for i, (a, b, c2) in enumerate(rows_data):
        bg = 'EBF3FB' if i % 2 == 0 else 'FFFFFF'
        for ci, v in enumerate([a, b, c2]):
            cell = tbl2.rows[i+1].cells[ci]
            cell.text = v
            set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)
    for ci, w in enumerate([1.5, 2.5, 2.5]):
        set_col_width(tbl2, ci, w)

    doc.add_paragraph()
    heading(doc, '4.5  The Abstract Syntax Tree for the Example Program', 2)
    body(doc,
        'After parsing, both parsers produce the same AST. Below is the '
        'tree for the first four statements (variable declarations):')
    doc.add_paragraph()
    code_block(doc, '''\
Program
  DeclStmt [int] age
    Literal [int] 20
  DeclStmt [float] gpa
    Literal [float] 3.75
  DeclStmt [string] name
    Literal [string] "Alice"
  DeclStmt [bool] passing
    Literal [bool] True''', font_size=9)

    body(doc, 'The if/else branch (Section 3 of the example) looks like this in the tree:')
    doc.add_paragraph()
    code_block(doc, '''\
  IfStmt
    condition:
      BinaryExpr ( >= )
        Identifier  gpa
        Literal [float] 3.5
    then:
      Block
        PrintStmt
          Literal [string] "Distinction"
    else:
      Block
        PrintStmt
          Literal [string] "Pass"''', font_size=9)

    body(doc, 'The while loop that sums 1 to 10:')
    doc.add_paragraph()
    code_block(doc, '''\
  WhileStmt
    condition:
      BinaryExpr ( <= )
        Identifier  i
        Literal [int] 10
    body:
      Block
        AssignStmt total =
          BinaryExpr ( + )
            Identifier  total
            Identifier  i
        AssignStmt i =
          BinaryExpr ( + )
            Identifier  i
            Literal [int] 1''', font_size=9)

    doc.add_page_break()

    # ══════════════════════════════════════════════════ 5. SEMANTIC ANALYSIS ═══
    heading(doc, '5.  Phase 3 – Semantic Analysis', 1)
    body(doc,
        'Semantic analysis walks the AST and enforces the "meaning" rules of '
        'the language — things a grammar cannot express: types must be compatible, '
        'variables must be declared before use, and names cannot be redeclared '
        'in the same scope. The SemanticAnalyzer class in semantic.py performs '
        'this phase.')

    doc.add_paragraph()
    heading(doc, '5.1  The Symbol Table', 2)
    body(doc,
        'The symbol table is the core data structure of semantic analysis. '
        'It maps each variable name to its Symbol record (name, type, '
        'declaration line). The SymbolTable class supports parent-chain lookup '
        'to implement block scoping:')
    doc.add_paragraph()
    two_col(doc, [
        ('define(sym)',      'Register a new symbol in the current scope.'),
        ('lookup(name)',     'Search current scope, then parent scopes (for use-before-declaration checking).'),
        ('lookup_local(name)', 'Search current scope only (for redeclaration checking — a variable CAN shadow an outer one).'),
        ('all_symbols()',    'Return every symbol visible at the current point (used to populate the Symbol Table tab in the IDE).'),
    ], header1='Method', header2='Purpose', w1=2.2, w2=4.3)

    doc.add_paragraph()
    body(doc,
        'A new SymbolTable is created each time a { } block is entered '
        '(_enter_scope), and the previous table is restored when the block ends '
        '(_exit_scope). This means variables declared inside an if or while block '
        'are invisible outside it.')

    doc.add_paragraph()
    heading(doc, '5.2  What the Analyser Checks', 2)
    two_col(doc, [
        ('Declaration before use',     'If a variable appears in an expression before it is declared, an ERROR is reported.'),
        ('No redeclaration',           'If the same name is declared twice in the same scope, an ERROR is reported.'),
        ('Type compatibility',         'The type of the initializer / right-hand side must be compatible with the declared variable type. '
                                        'int → float widening is allowed; other cross-type assignments are errors.'),
        ('Arithmetic operands',        'Operators +, -, *, /, % require both operands to be int or float (except + which also allows string + string).'),
        ('Comparison types',           '==, !=, <, >, <=, >= require both operands to be the same type (or both numeric).'),
        ('Condition types (warning)',   'if and while conditions that are not bool produce a WARNING (not an error).'),
        ('Logical operands (warning)',  '&&, || applied to non-bool operands generate a WARNING.'),
    ], header1='Rule', header2='Detail', w1=2.3, w2=4.2)

    doc.add_paragraph()
    heading(doc, '5.3  Walking Through the Example Program', 2)
    body(doc,
        'The analyser visits every node in the AST. Here is the trace for the '
        'most important statements:')
    doc.add_paragraph()

    three_col(doc, [
        ('int age = 20;',
         'DeclStmt',
         'Look up "age" locally → not found (OK). Infer type of 20 → int. int compatible with int. '
         'Define Symbol("age", "int", line 2). ✓'),
        ('float gpa = 3.75;',
         'DeclStmt',
         'Look up "gpa" locally → not found. Infer 3.75 → float. float compatible with float. '
         'Define Symbol("gpa", "float", line 3). ✓'),
        ('string name = "Alice";',
         'DeclStmt',
         'Infer "Alice" → string. string compatible with string. Define Symbol("name","string",line 4). ✓'),
        ('bool passing = true;',
         'DeclStmt',
         'Infer true → bool. bool compatible with bool. Define Symbol("passing","bool",line 5). ✓'),
        ('int sum = a + b;',
         'DeclStmt + BinaryExpr',
         'lookup("a") → int. lookup("b") → int. '
         'int + int → result type int. int compatible with int. Define sum:int. ✓'),
        ('if (gpa >= 3.5)',
         'IfStmt',
         'Infer gpa → float (from symbol table). Infer 3.5 → float. '
         'float >= float → result type bool. Bool condition ✓. Enter new scope for then block. ✓'),
        ('bool honours = passing && gpa >= 3.5;',
         'DeclStmt + BinaryExpr',
         'gpa >= 3.5 → bool. passing → bool. bool && bool → bool. '
         'bool compatible with bool. Define honours:bool. ✓'),
        ('while (i <= 10)',
         'WhileStmt',
         'lookup("i") → int. 10 → int. int <= int → bool. '
         'Bool condition ✓. Enter new scope for body block.'),
        ('total = total + i;',
         'AssignStmt',
         'lookup("total") → int. total + i → int + int → int. '
         'int compatible with int target. ✓'),
    ],
    headers=('Statement', 'Node Type', 'Semantic Analyser Action'),
    widths=(2.0, 1.3, 3.2))

    doc.add_paragraph()
    heading(doc, '5.4  Symbol Table After Full Analysis', 2)
    body(doc,
        'After the analyser finishes, the global symbol table contains all '
        'top-level variables. (Variables declared inside blocks exist only for '
        'the duration of that block\'s scope and are not visible globally.)')
    doc.add_paragraph()

    three_col(doc, [
        ('age',      'int',    'line 2'),
        ('gpa',      'float',  'line 3'),
        ('name',     'string', 'line 4'),
        ('passing',  'bool',   'line 5'),
        ('a',        'int',    'line 14'),
        ('b',        'int',    'line 15'),
        ('sum',      'int',    'line 16'),
        ('diff',     'int',    'line 17'),
        ('prod',     'int',    'line 18'),
        ('modulo',   'int',    'line 19'),
        ('honours',  'bool',   'line 34'),
        ('total',    'int',    'line 40'),
        ('i',        'int',    'line 41'),
        ('n',        'int',    'line 51'),
        ('farewell', 'string', 'line 62'),
    ],
    headers=('Variable', 'Type', 'Declared at'),
    widths=(1.6, 1.6, 3.3))

    doc.add_paragraph()
    heading(doc, '5.5  Type Compatibility Rules', 2)
    body(doc,
        'The analyser allows a limited form of widening: an int value can be '
        'assigned to a float variable (because every integer is representable '
        'as a float). All other cross-type assignments are errors:')
    doc.add_paragraph()
    three_col(doc, [
        ('int   → int',    '✓ OK',  'Same type'),
        ('int   → float',  '✓ OK',  'Widening: every int fits in float'),
        ('float → int',    '✗ Error','Loss of precision not allowed'),
        ('string → int',   '✗ Error','Incompatible types'),
        ('bool  → int',    '✗ Error','Incompatible types'),
        ('string + string','✓ OK',  'Special case: + is string concatenation'),
    ],
    headers=('Assignment / Operation', 'Result', 'Reason'),
    widths=(2.0, 1.0, 3.5))

    doc.add_page_break()

    # ══════════════════════════════════════════════════ 6. ERROR EXAMPLES ═════
    heading(doc, '6.  What Happens When Code Contains Errors', 1)
    body(doc,
        'The compiler detects errors in each phase and stops at the phase '
        'where the first error occurs. The web IDE displays errors in red '
        'in the corresponding tab (Tokens for lex errors, Parse Steps for '
        'syntax errors, Semantic for semantic errors).')

    doc.add_paragraph()
    heading(doc, '6.1  Lexical Error', 2)
    code_block(doc, 'int x = 10@;')
    body(doc, '→ Error: Line 1, Col 10: Unknown character \'@\'')
    body(doc, 'The "@" character is not in any token pattern; the Lexer logs the error and skips it.')

    doc.add_paragraph()
    heading(doc, '6.2  Syntax Error', 2)
    code_block(doc, 'int x = ;')
    body(doc, '→ Error: Line 1: Expected primary expression, got SEMICOLON (\';\')' )
    body(doc,
        'After "=" the parser expects an expression. It finds ";" instead. '
        'Panic-mode recovery skips tokens until the next ";" or "}" is found.')

    doc.add_paragraph()
    heading(doc, '6.3  Semantic Error — Use Before Declaration', 2)
    code_block(doc, 'int y = x + 1;   // x not yet declared')
    body(doc, '→ Error: Line 1: Undefined variable \'x\'')
    body(doc,
        'The analyser calls lookup("x") which searches the current scope '
        'chain and finds nothing. An error is recorded and analysis continues.')

    doc.add_paragraph()
    heading(doc, '6.4  Semantic Error — Type Mismatch', 2)
    code_block(doc, 'int z = "hello";')
    body(doc, '→ Error: Line 1: Cannot assign string to int variable \'z\'')
    body(doc,
        'The analyser infers the type of "hello" as string, then calls '
        '_compatible("int", "string") which returns False.')

    doc.add_paragraph()
    heading(doc, '6.5  Semantic Error — Redeclaration', 2)
    code_block(doc, 'int x = 1;\nint x = 2;   // same name in same scope')
    body(doc, '→ Error: Line 2: Variable \'x\' already declared in this scope')
    body(doc,
        'lookup_local("x") finds x in the current scope (not just parent scopes), '
        'so redeclaration is rejected.')

    doc.add_paragraph()
    heading(doc, '6.6  Semantic Warning — Non-Bool Condition', 2)
    code_block(doc, 'int flag = 1;\nif (flag) { print("yes"); }')
    body(doc, '→ Warning: Line 2: if-condition is \'int\', expected \'bool\'')
    body(doc,
        'This is a warning, not an error — the program still compiles and runs. '
        'The condition is treated as truthy (any non-zero int). '
        'Using a proper bool condition is recommended.')

    doc.add_page_break()

    # ══════════════════════════════════════════════════ 7. HOW TO RUN ═══════════
    heading(doc, '7.  How to Run the Compiler', 1)

    heading(doc, '7.1  Web Interface (Recommended)', 2)
    code_block(doc, 'cd /home/joselyn/Desktop/complier\npython3 app.py')
    body(doc, 'Then open  http://localhost:5000  in a browser.')
    doc.add_paragraph()
    two_col(doc, [
        ('Example button',         'Loads the full demo program into the editor.'),
        ('▶  Compile button',      'Runs Phases 1–3 (Lex → Parse → Semantic). No execution.'),
        ('⚡  Run Code button',    'Runs all four phases including execution; output appears in the Console tab.'),
        ('Parser selector',        'Switch between Top-Down and Bottom-Up to compare their parse-step logs.'),
        ('Tokens tab',             'Shows every token with its type, value, line, and column.'),
        ('Parse Steps tab',        'Full shift/reduce or predict/expand trace.'),
        ('AST tab',                'Pretty-printed Abstract Syntax Tree.'),
        ('Semantic tab',           'Analyser log with scope entry/exit and any errors or warnings.'),
        ('Symbol Table tab',       'All declared variables with their types and declaration lines.'),
        ('Console tab',            'Execution output (print results) and the execution trace log.'),
    ], header1='Control / Tab', header2='Purpose', w1=2.3, w2=4.2)

    doc.add_paragraph()
    heading(doc, '7.2  Desktop GUI', 2)
    code_block(doc, 'python3 gui.py')
    body(doc,
        'The Tkinter GUI provides the same functionality in a dark-themed '
        'desktop window with a built-in code editor and tabbed output panels.')

    doc.add_paragraph()
    heading(doc, '7.3  Python API', 2)
    code_block(doc, '''\
from compiler import Compiler
c = Compiler()
result = c.compile(source_code, parser_type='topdown', execute=True)
print(result['exe_output'])   # list of print() output lines
print(result['symbol_table']) # dict of name → "type (line N)"''')

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p.add_run('— End of Document —')
    r2.italic = True
    r2.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    # ── save ──────────────────────────────────────────────────────────────────
    out = '/home/joselyn/Desktop/complier/Compiler_Explanation.docx'
    doc.save(out)
    print(f'Saved → {out}')


if __name__ == '__main__':
    build()
